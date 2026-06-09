# -*- coding: utf-8 -*-
"""템플릿(232320635_하현성_보고서.docx)을 그대로 베이스로, 본문만 FF 내용으로 교체.
표지/표지표/목차/TOC/스타일 전부 보존. 존댓말, 시행착오 상세, 역할분담, 공유산출물 반영."""
import os, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TPL = '참고자료/232320635_하현성_보고서.docx'
SUB = '제출_딥러닝응용_3조_김주헌_하현성'
FIG = '_ppt_assets_iter6'
OUT = os.path.join(SUB, '232320635_하현성_딥러닝응용_기말보고서.docx')

shutil.copy2(TPL, OUT)
doc = Document(OUT)
body = doc.element.body
WT = qn('w:t')

# ---------- 1) 표지 텍스트 교체 ----------
def replace_text_everywhere(old, new):
    for t in body.iter(WT):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            return True
    return False

def set_para_text(contains, newtext):
    for p in doc.paragraphs:
        if contains in p.text:
            if p.runs:
                p.runs[0].text = newtext
                for r in p.runs[1:]:
                    r.text = ''
            else:
                p.add_run(newtext)
            return True
    return False

replace_text_everywhere('중간고사 보고서', '기말 프로젝트 보고서')
set_para_text('Midterm take-home Exam', '[ Final Project ]  멀티모달 딥러닝 기반 제조 공정 이상 탐지')
# 표지표 이름: 하현성 -> 하현성 · 김주헌 (3조)
for t in body.iter(WT):
    if t.text and t.text.strip().endswith('하현성') and '이름' not in t.text:
        t.text = t.text.replace('하현성', '하현성 · 김주헌 (3조)')

# ---------- 2) 본문(첫 Heading1 '과제 개요 및 실험 환경'부터) 제거, sectPr 보존 ----------
children = list(body)
start = None
for i, ch in enumerate(children):
    txt = ''.join((t.text or '') for t in ch.iter(WT))
    if ch.tag == qn('w:p') and '과제 개요 및 실험 환경' in txt:
        start = i
        break
assert start is not None, '본문 시작점 못찾음'
for ch in children[start:]:
    if ch.tag == qn('w:sectPr'):
        continue
    body.remove(ch)

# ---------- 3) 헬퍼 (모든 신규 단락은 sectPr 앞에 자동 삽입) ----------
NAVY = RGBColor(0x1E, 0x2A, 0x4A)

import re
def _strip_num(text):
    # 제목 스타일이 자동 번호를 매기므로 수동 번호("1.", "2.1 ") 제거
    return re.sub(r'^\d+(\.\d+)*\.?\s+', '', text)

def H1(text):
    p = doc.add_paragraph(_strip_num(text), style='Heading 1')
    p.paragraph_format.space_before = Pt(18)   # 다음 섹션 앞은 넉넉히
    p.paragraph_format.space_after = Pt(3)      # 제목과 본문은 붙게
    return p

def H2(text):
    p = doc.add_paragraph(_strip_num(text), style='Heading 2')
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(2)
    return p

def body_p(text):
    p = doc.add_paragraph(text, style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('•  ' + text)
    return p

def code_block(text):
    # 템플릿의 '코드 블럭' 스타일 사용
    try:
        p = doc.add_paragraph(text, style='코드 블럭')
    except KeyError:
        p = doc.add_paragraph(text, style='Normal')
        for r in p.runs:
            r.font.name = 'Consolas'; r.font.size = Pt(9)
    return p

def figure(fname, width_cm, caption):
    path = os.path.join(FIG, fname)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    cp = doc.add_paragraph(caption, style='Normal'); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cp.runs:
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = 'Table Grid'
    except KeyError:
        pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = c._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:val'):'clear', qn('w:fill'):'DCE6F1'})
        c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ''
            run = cells[j].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ============================================================
# 1. 요약
# ============================================================
H1('1. 요약')
body_p('본 프로젝트는 실제 제조 조립 라인에서 수집된 시계열 센서 데이터와 카메라 이미지를 함께 활용하여, 제품 조립 결함을 매 측정 시점 단위로 분류하는 멀티모달 이상 탐지 모델을 구현하고 단일 모달리티 모델과 정량적으로 비교한 연구입니다. 사우스캐롤라이나대학교 Future Factories(FF) 공개 데이터셋(166,001 시점, 285 사이클, 5개 상태)을 사용하였으며, 센서는 양방향 LSTM(BiLSTM)으로, 이미지는 ImageNet 사전학습 ResNet18 전이학습으로 모델링한 뒤, 세 가지 융합 전략(Decision-level, Concat, Cross-Attention)을 비교하였습니다.')
body_p('프로젝트 초기(중간발표 시점)에는 이미지 데이터가 부족하여 융합 모델이 센서 단독보다 오히려 낮은 성능을 보였습니다. 이를 데이터 불균형 문제로 진단하고, 이미지 데이터 확장·센서 및 이미지 브랜치 가중치 이식(warm-start)·확률 결합(Decision-level)·학습률 안정화 등 여러 단계의 개선을 거쳤습니다. 그 결과 두 모달리티가 모두 존재하는 구간(3,912 시점)에서 특징을 학습형으로 결합한 Concat 융합이 가중 F1 0.977로 센서 단독(0.907)과 이미지 단독(0.957)을 모두 앞섰으며, 그 차이는 부트스트랩 검정에서 통계적으로 유의하였습니다.')
body_p('나아가 학습된 모델을 고정한 채 테스트 이미지만 단계적으로 열화시키는 추가 실험을 수행한 결과, 이미지 품질이 나빠질수록 고정 가중 융합은 취약해지고, 이미지 신뢰도에 따라 비중을 자동으로 조절하는 적응형(Cross-Attention) 융합이 가장 견고함을 확인하였습니다. 결론적으로 멀티모달 융합의 가치는 무조건적인 정확도 향상이 아니라 입력 품질이 나쁜 조건에서의 견고성에 있으며, 어떤 융합 방식을 선택하는지가 결정적이라는 점을 정량적으로 규명하였습니다.')

# ============================================================
# 2. 서론
# ============================================================
H1('2. 서론')
H2('2.1 주제 선택 동기')
body_p('본 프로젝트의 출발점은 팀원들의 현장 경험입니다. 팀원 전원이 자동화 장비 업계에 종사하면서 제조 공정에서 딥러닝 기반 불량 판정이 어떻게 활용되는지를 직접 접해 왔습니다. 현재 산업 현장의 불량 판정은 대부분 비전 카메라 단독 방식이지만, 실제로는 조명·각도·가려짐 때문에 비전만으로 판단이 어려운 경우가 자주 발생합니다. 반대로 센서 데이터 역시 노이즈나 고장 때문에 단독으로는 신뢰하기 어려운 경우가 존재합니다.')
body_p('이미지는 시각적 결손을 잘 잡아내지만 가려짐·조명 변화에 취약하고, 센서는 동작 패턴은 정확히 포착하지만 시각 정보가 부재합니다. 두 모달리티의 약점이 서로 보완될 수 있다는 점에 주목하여, 이미지와 센서를 함께 활용하는 멀티모달 접근이 실제로 어떤 성능 차이를 만드는지를 정량적으로 검증하는 것을 핵심 동기로 삼았습니다.')
H2('2.2 프로젝트 목표 및 접근')
body_p('구체적인 목표는 동일한 데이터와 동일한 평가 지표 위에서 (1) 센서 단독 모델, (2) 이미지 단독 모델, (3) 멀티모달 융합 모델을 공정하게 비교하여, 멀티모달 융합이 단일 모달리티 대비 실제로 어떤 이점을 주는지를 직접 구현하고 분석하는 것입니다. 융합은 강의에서 학습한 개념(전이학습, RNN/LSTM, 멀티모달 결합)을 토대로 Decision-level(확률 가중 평균), Concat(특징 결합 후 공동 학습), Cross-Attention(센서 질의가 이미지에 주의를 거는 적응형)의 세 가지로 구현하였으며, 이후 이미지 열화 견고성 실험을 통해 멀티모달의 실질적 이점을 검증하였습니다.')
H2('2.3 관련 연구 — NSF-MAP (IJCAI 2025)')
body_p('주요 참고 논문은 동일한 Future Factories 데이터셋을 사용한 NSF-MAP(IJCAI 2025)입니다. 이 연구는 단계적 발전 전략을 취하여, 기본 Decision-Level Fusion(P1)에서 정확도 72%, 전이학습을 더한 P2에서 88%, Process Ontology 기반 Knowledge Infusion을 더한 P3에서 93%를 달성합니다. 본 프로젝트는 NSF-MAP을 참고하되, 시계열은 Autoencoder 대신 강의에서 학습한 BiLSTM으로, 이미지는 EfficientNet-B0 대신 ResNet18로 재구성하여 강의 학습 범위 안에서 멀티모달 융합의 효과 자체를 검증하는 데 집중하였습니다. Knowledge Infusion·Ontology는 여건상 향후 확장 과제로 분류하였습니다.')

# ============================================================
# 3. 데이터
# ============================================================
H1('3. 데이터')
H2('3.1 데이터셋 출처 · 링크 · 확보 과정')
body_p('사우스캐롤라이나대학교 Future Factories(FF) Lab이 공개한 산업용 멀티모달 데이터셋을 사용하였습니다. 실제 조립 라인을 모사한 시설에서 산업용 로봇 4대와 컨베이어가 약 30시간 동안 로켓 모형을 조립하며 약 1.95Hz로 측정하여, 총 166,001개의 시점(레코드)과 285개 사이클(제품)이 기록되었습니다. 데이터는 (1) 40개 이상의 센서 채널과 이상 라벨·사이클 정보·이미지 경로가 담긴 시계열 CSV와 (2) 카메라 2대가 동기 촬영한 이미지로 구성됩니다. 결함은 의도적으로 특정 부품을 빼고 조립하여 생성하였고 누적적이며(코→몸통2→몸통1 순), 각 시점에 그 순간의 상태가 라벨로 기록되어 있습니다. 본 프로젝트에서는 소수 클래스를 일부 병합하여 5개 클래스로 단순화하였습니다.')
body_p('데이터 확보는 두 갈래로 진행하였습니다. 센서·라벨 데이터는 NSF-MAP 논문 저자의 공개 저장소(GitHub: ChathurangiShyalika/NSF-MAP)에 연결된 Google Drive 링크에서 전처리 CSV(FF_Multimodal.csv, 약 75MB, 166,001행 × 51열)를 직접 내려받아 사용하였고, 이후 Kaggle 노트북에서 활용하기 위해 동일 CSV를 Kaggle 데이터셋으로 재업로드(공개)하였습니다. 반면 이미지 원본은 Kaggle에 6개 파트(각 약 104GB, 합계 약 580GB)로 분할 업로드되어 있어 로컬 PC로 전부 내려받는 것이 사실상 불가능했습니다. 따라서 전체 다운로드 대신 Kaggle 노트북의 Add Input 기능으로 실제 필요한 2개 파트(part 16, part 26)만 세션에 마운트하여, 다운로드 없이 서버에 적재된 이미지를 직접 읽어 사용하였습니다.')
table(['구성', '내용', '용량', '출처 · 확보 방법'],
      [['센서·라벨 CSV', '166,001행 × 51열 (센서 40+채널, 라벨, 사이클, 이미지 경로)', '약 75MB', 'NSF-MAP GitHub→Google Drive 다운로드 (Kaggle 데이터셋으로 재업로드)'],
       ['이미지 원본', '카메라 2대 동기 촬영 PNG', '6파트×~104GB ≈ 580GB', 'Kaggle ramyharik 데이터셋 (필요 2파트만 Add Input 마운트)']])
bullet('센서 CSV(재업로드, 공개): https://www.kaggle.com/datasets/hyunsungha/ff-multimodal-csv')
bullet('센서 CSV 원 출처(NSF-MAP): https://github.com/ChathurangiShyalika/NSF-MAP')
bullet('이미지 데이터셋 part 16: https://www.kaggle.com/datasets/ramyharik/ff-2023-12-12-multi-modal-dataset-16')
bullet('이미지 데이터셋 part 26: https://www.kaggle.com/datasets/ramyharik/ff-2023-12-12-multi-modal-dataset-26')
table(['상태(클래스)', '개수', '비중'],
      [['Normal', '90,775', '54.7%'],
       ['NoNose, NoBody2, NoBody1 (병합)', '26,628', '16.0%'],
       ['NoNose, NoBody2', '25,206', '15.2%'],
       ['NoNose', '19,307', '11.7%'],
       ['NoBody1 (병합)', '4,016', '2.4%']])
H2('3.2 데이터 특성')
body_p('센서 데이터는 원본 CSV의 40개 이상 채널 중, 조립 동작·결함과 직접 관련된 22개 채널을 선별하여 사용하였습니다. 선별한 22개 채널은 로봇 4대의 그리퍼 하중·위치, 모터 구동기 온도, 컨베이어 속도, 로봇 관절 각도로 구성되며, 모두 모든 시점에 결측 없이 기록되어 시계열 모델(BiLSTM)의 안정적인 입력이 됩니다. 부품이 누락되면 그리퍼가 부품을 집을 때의 하중과 관절 동작 패턴이 달라지므로, 이러한 센서 신호만으로도 상태를 상당 부분 구분할 수 있습니다.')
table(['채널 그룹 (예시 컬럼명)', '개수', '설명'],
      [['그리퍼 로드셀 (I_R0x_Gripper_Load)', '4', '로봇 4대 그리퍼가 부품을 집을 때의 하중 — 부품 누락 시 패턴 변화'],
       ['그리퍼 포텐셔미터 (I_R0x_Gripper_Pot)', '4', '그리퍼 개폐 위치(변위)'],
       ['VFD 온도 (Q_VFDx_Temperature)', '4', '모터 구동기(가변주파수 드라이브) 온도'],
       ['컨베이어 속도 (M_Convx_Speed_mmps)', '4', '컨베이어 4개 구간 이송 속도(mm/s)'],
       ['로봇 관절각 (M_R01/R04 S·L·U JointAngle)', '6', '로봇 1·4의 베이스(S)·하부(L)·상부(U) 관절 각도']])
bullet('이미지: 카메라 2대가 한 시점에 2장을 동기 촬영합니다. 부품이 명확히 보이는 공정 단계(CycleState 4·9)에서만 사용하며, 두 카메라가 모두 매칭되는 이미지 시점은 19,234개입니다.')
bullet('클래스 불균형(Normal 과반)이 크므로, 평가지표로는 가중 F1(Weighted F1)을 채택하였습니다.')
H2('3.3 전처리 및 학습/평가 분리')
body_p('센서 22채널을 표준화(StandardScaler)하고 라벨을 정수 인코딩한 뒤, 시계열은 길이 50의 슬라이딩 윈도우로 구성하여 윈도우의 끝 시점 상태를 예측하도록 하였습니다. 이미지는 224×224로 리사이즈·정규화하였습니다. 가장 중요한 점은 데이터 누수를 방지하기 위해 행(시점) 단위가 아니라 사이클 단위로 학습/평가를 8:2로 분리하였다는 것입니다. 같은 사이클의 시점들이 학습과 평가에 동시에 들어가지 않으므로 인접 시점의 정보가 새어 들어가는 누수를 차단하며, 분할 시 사이클별 결함 유무로 층화(stratify)하여 클래스 분포를 보존하였습니다.')
code_block("from sklearn.model_selection import train_test_split\n"
           "# 사이클 단위 8:2 분할 — 같은 사이클이 train/test에 섞이지 않게(데이터 누수 방지)\n"
           "train_cycles, test_cycles = train_test_split(\n"
           "    cycles, train_size=TRAIN_RATIO, random_state=SEED,\n"
           "    stratify=[cycle_has_anomaly[c] for c in cycles])  # 결함 유무로 층화")
figure('code_split.png', 15.5, '[그림 1] 사이클 단위 학습/평가 분리 코드 (Kaggle 노트북)')
body_p('평가는 목적에 따라 세 가지 집합으로 구분하였습니다. ① 센서 전체(모든 CycleState) 30,132 시점, ② 이미지가 사용되는 cycle 4·9 전체 9,521 시점(이미지 결측 약 59% 포함, 공정 비교용 FAIR 셋), ③ 그중 이미지가 실제 존재하는 3,912 시점(41%, 단독과 융합을 정면으로 비교할 수 있는 유일한 구간)입니다.')
H2('3.4 Kaggle 환경 구성 및 경로 매핑 (시행착오)')
body_p('앞서 기술한 대용량 이미지 문제(약 580GB) 때문에 학습은 Kaggle Notebook 환경에서 진행하였습니다. Kaggle은 서버에 마운트된 데이터셋을 다운로드 없이 직접 읽을 수 있고 무료 Tesla T4 GPU(주 30시간)를 제공하므로, 데이터 용량과 연산 자원 제약을 동시에 우회할 수 있었습니다. 다만 학습용 GPU가 팀원 로컬 PC에 부재하여 이 무료 GPU 시간이 곧 자원 제약이 되었습니다.')
body_p('이 환경을 구성하는 과정에서 다음과 같은 문제들을 차례로 해결하였습니다. (1) ResNet18의 ImageNet 사전학습 가중치 다운로드가 네트워크 차단으로 실패하여, Kaggle 계정의 휴대폰 인증 후 Internet 옵션을 활성화하여 해결하였습니다. (2) CSV 경로와 이미지 BATCH 폴더 위치가 환경마다 달라, 하드코딩 대신 glob 기반 자동 탐색으로 전환하였습니다. (3) CSV의 이미지 상대경로에 포함된 "Dataset/" 접두사가 실제 저장 구조와 불일치하여 매칭이 0건이던 문제를, 접두사를 제거하는 경로 정규화로 해결하였습니다.')
code_block("def find_image_path(rel):                 # 'Dataset/' 접두사 제거 + 자동 탐색\n"
           "    if rel.startswith('Dataset/'): rel = rel[len('Dataset/'):]\n"
           "    for base in IMAGE_BASE_DIRS:\n"
           "        cand = os.path.join(base, rel)\n"
           "        if os.path.exists(cand): return cand\n"
           "    return None")

# ============================================================
# 4. 방법
# ============================================================
H1('4. 방법')
H2('4.1 비교 설계')
body_p('동일한 데이터와 동일한 사이클 분할 위에서 세 단독/융합 모델을 공정하게 비교합니다. 멀티모달의 효과를 분리해서 보기 위해, 이미지가 실제로 존재하는 구간(3,912 시점)에서 단독과 융합을 직접 비교하고, 이미지 결측이 포함된 전체 구간(9,521 시점)에서도 함께 확인하였습니다.')
H2('4.2 단독 모델')
bullet('센서(BiLSTM): 2층 양방향 LSTM(은닉 128)으로 22채널·길이 50 시퀀스를 인코딩하여 분류합니다. 강의의 RNN/LSTM 내용을 적용하였습니다.')
bullet('이미지(ResNet18): torchvision의 ImageNet 사전학습 ResNet18을 전이학습하며, 두 카메라의 예측을 평균합니다. 학습률은 1e-3에서 학습이 불안정하여 1e-4로 낮춰 안정화하였습니다(전체 미세조정에는 1e-3이 과대).')
body_p('오픈소스 출처와 수정 사항: 이미지 백본은 PyTorch torchvision의 ResNet18(ImageNet 사전학습 가중치)을 사용하였습니다. 본 과제에 맞추어 최종 분류층(fc)을 5개 클래스로 교체하고, 두 카메라 입력의 예측/특징을 평균하도록 멀티뷰 처리를 추가하였습니다.')
H2('4.3 융합 모델 (세 가지)')
bullet('Decision-level(고정 가중): 센서·이미지 각 모델의 softmax 확률을 고정 비율 w로 가중 평균합니다. 이미지가 없으면 센서로 자동 폴백합니다. 단순하고 빠르지만 가중치가 고정이라 이미지가 나빠져도 비중을 줄이지 못합니다.')
bullet('Concat(학습형): 센서 특징(256)과 이미지 특징(512)을 이어붙여(768) 공동 FC로 학습하여 두 정보를 비선형 결합합니다.')
bullet('Cross-Attention(적응형): 센서 특징을 질의(query)로, 이미지의 공간 토큰(7×7=49개, 두 카메라면 98개)을 키/값으로 두어 어텐션합니다. 이미지 신뢰도가 낮으면 어텐션 기여가 작아져 센서 경로가 분류를 지배하므로 자동 폴백이 이루어집니다.')
figure('code_crossattn.png', 16.0, '[그림 2] Cross-Attention 융합 forward — 이미지 신뢰도에 따라 비중 자동 조절 (Kaggle 노트북)')
body_p('세 융합 모두 단독 모델의 사전학습 가중치를 이식(warm-start)한 뒤 학습하였습니다. 강의에서 다룬 멀티모달 결합 개념을 토대로 세 가지 결합 수준(결정·특징·주의)을 직접 구현·비교한 점이 본 프로젝트의 핵심입니다.')
H2('4.4 구현 중 발생한 문제와 해결 (디버깅)')
body_p('이미지 모델 학습 초기에, 검증 F1이 여러 epoch 동안 전혀 변하지 않고 모든 샘플을 다수 클래스로 분류하는 편향 학습이 관찰되었습니다. 원인은 클래스 가중치(class weight)를 센서 시퀀스 데이터(약 12만 개)의 라벨 분포로 계산한 뒤, 분포가 전혀 다른 이미지 데이터(약 9,500장) 학습에 그대로 적용한 데 있었습니다. 이미지 데이터셋 전용 클래스 가중치를 별도로 계산하고, 학습 함수에 모달리티별 가중치를 주입할 수 있도록 일반화하여 해결하였습니다. 또한 테스트 셋에 일부 클래스만 등장하여 발생한 classification_report의 클래스 개수 불일치 오류는 labels 인자를 명시하여 해결하였습니다. 이 경험을 통해 멀티모달 학습에서는 모달리티별 데이터 분포 차이를 반드시 고려해야 한다는 점을 배웠습니다.')

# ============================================================
# 5. 실험결과
# ============================================================
H1('5. 실험 결과')
H2('5.1 중간발표 시점의 결과와 진단')
body_p('초기 구현(이미지 1개 파트, 약 9,542장, 단순 Concat)에서는 센서 단독이 가중 F1 0.927로 가장 높고, 이미지 단독은 0.611, 융합은 0.882로, 융합이 센서 단독보다 오히려 낮은 예상 밖의 결과가 나왔습니다. 원인을 분석한 결과, 이미지 데이터는 Cycle 4·9의 한 파트만 사용하여 양이 적고 일부 클래스만 등장하는 반면 센서는 전체 사이클의 약 12만 시퀀스로 훨씬 풍부하여, 융합 시 이미지 특징이 정보를 더하기보다 노이즈처럼 작용해 센서 정보를 희석시킨 것으로 판단하였습니다. 이는 NSF-MAP 논문이 기본 융합(P1)에서 72%로 가장 낮았다가 단계적 개선으로 향상된 것과 동일한 맥락으로, 멀티모달 융합에서 데이터 균형이 핵심임을 시사합니다.')
H2('5.2 개선 과정')
bullet('평가셋 정렬: 센서·이미지 단독은 전체 사이클에서, 융합은 cycle 4·9 부분집합에서 평가되어 직접 비교가 부당했던 문제를 발견하고, 모든 모델을 동일한 cycle 4·9 셋에서 재평가하였습니다(센서는 이 구간에서 0.871로, 전체 0.922보다 낮아 더 어려운 구간임을 확인).')
bullet('브랜치 가중치 이식(warm-start): 융합의 센서·이미지 브랜치를 각각 학습된 단독 모델 가중치로 초기화하여, 소량 데이터로 처음부터 학습할 때의 약화를 해소하였습니다.')
bullet('Decision-level(확률) 융합 도입: 학습형 융합 헤드가 이미지 결측 샘플 탓에 센서 의존적으로 수렴하는 한계를 우회하여, 두 강한 단독 모델의 확률을 결합하고 가중치 w는 학습 셋에서 선택(테스트 누수 방지)하였습니다.')
bullet('이미지 데이터 확장 및 학습률 안정화: 이미지 파트를 늘려 커버리지를 높이고, 이미지 학습률을 1e-4로 조정하여 학습을 안정화함으로써 이미지 단독 성능을 크게 회복시켰습니다.')
H2('5.3 최종 결과 — 단독 vs 융합 (이미지 존재 구간 3,912 시점)')
body_p('두 모달리티가 모두 존재하는 구간에서, 특징을 학습형으로 결합한 Concat 융합이 0.977로 두 단독을 모두 앞섰습니다. 이미지 단독이 이미 0.957로 강하여 절대 이득은 작지만(+0.020), 융합과 센서의 차이는 부트스트랩(1000회 재추출) 검정에서 평균 +0.050, 95% 신뢰구간 [+0.039, +0.060]으로 0보다 커 통계적으로 유의하였습니다.')
table(['모델', '센서', '이미지', 'Concat', 'Cross-Attn', 'Decision'],
      [['가중 F1', '0.907', '0.957', '0.977', '0.906', '0.957']])
figure('res_present.png', 13.5, '[그림 3] 이미지 존재 구간(3,912 시점) 단독 vs 융합 가중 F1')
H2('5.4 결측 포함 전체 구간 (cycle 4·9, 9,521 시점)')
body_p('이미지가 없는 시점까지 포함한 전체 구간에서도 융합(Concat 0.900, Decision 0.893)이 센서 0.871·이미지 0.826을 모두 상회하여, 이미지 결측이 많은 현실적 분포에서도 융합이 단독보다 견고함을 확인하였습니다. 참고로 센서 단독을 모든 CycleState(30,132 시점)에서 평가하면 0.9215였습니다.')
table(['모델', '센서', '이미지', 'Concat', 'Cross-Attn', 'Decision'],
      [['가중 F1', '0.871', '0.826', '0.900', '0.871', '0.893']])
H2('5.5 이미지 열화 견고성 (추가 실험)')
body_p('학습된 모델을 고정한 채 테스트 이미지만 깨끗→중간열화→강열화 3단계로 약화시켜(블러+조명변화+가림) 추론만 비교하였습니다. 이미지 단독이 0.957→0.681→0.028로 붕괴하는 동안 융합 방식마다 다른 결과를 보였습니다. 고정 가중 Decision은 중간 열화만으로 센서(0.907) 아래로 떨어졌고(0.775), 학습형 Concat은 중간까지 버틴 뒤(0.931) 강열화에 붕괴(0.303)하였으며, 적응형 Cross-Attention만 끝까지 견고(0.905)하였습니다. 다만 강열화(이미지 0.03)는 거의 파괴에 가까운 극단 가정으로 경향 확인용입니다.')
table(['이미지 품질', 'Image', 'Concat', 'Decision', 'Cross-Attn', 'Sensor'],
      [['깨끗', '0.957', '0.977', '0.957', '0.906', '0.907'],
       ['중간열화', '0.681', '0.931', '0.775', '0.905', '0.907'],
       ['강열화', '0.028', '0.303', '0.116', '0.905', '0.907']])
figure('res_robust.png', 12.5, '[그림 4] 이미지 열화 단계별 모델 견고성')
H2('5.6 결함 유형별 성능')
body_p('Concat 융합 기준으로 단일 결함(NoNose)은 0.97로 정확하나, 3부품 복합 결함은 0.71로 가장 어려웠습니다. 3부품 결함은 2부품 결함에 부품 하나가 더 빠진 형태라 시각·센서적으로 닮아 인접 클래스와 혼동되기 때문입니다. 단순 정상/불량 이진 판정이라면 쉬웠을 문제로, 난이도는 결함 유형까지 세분류하는 데서 비롯합니다.')
figure('res_byclass.png', 13.0, '[그림 5] 결함 유형별 F1 (Concat, cycle 4·9)')

# ============================================================
# 6. 토의 내용
# ============================================================
H1('6. 토의 내용')
H2('6.1 배운 점과 해석')
bullet('멀티모달의 가치는 무조건적 정확도 향상이 아니라, 입력 품질이 나쁜 조건에서의 견고성에 있습니다.')
bullet('같은 멀티모달이라도 융합 방식이 결정적입니다 — 정확도는 Concat이 전 구간에서 최고, 열화 견고성은 적응형(Cross-Attention)만 확보하였습니다.')
bullet('멀티모달 융합의 이득은 두 모달리티의 데이터 균형에 크게 의존하며, 모달리티별 데이터 분포 차이(클래스 가중치 등)를 반드시 고려해야 합니다.')
bullet('전이학습에서 학습률이 결과를 좌우함을 체감하였고(이미지 LR 1e-3→1e-4로 안정화), 사이클 단위 분할 같은 평가 설계가 수치의 신뢰성에 직접 영향을 준다는 점을 배웠습니다.')
H2('6.2 한계 및 향후 계획')
bullet('(중요) 본 1차 제출에서는 Kaggle GPU 자원 한도로 인해, 학습/평가를 사이클 단위로 분리하되 별도의 검증(validation) 셋을 추가로 분리하지 못하고 모델 선택(early stopping)에 평가셋 지표를 사용하였습니다. 이는 엄밀성 측면의 한계이며, 교수님 피드백 반영 2차 제출 시 학습 사이클에서 검증 사이클을 별도로 분리한 train/validation/test 3분할로 재실행하여 반영할 계획입니다.')
bullet('단일 데이터셋 검증이므로 일반화를 위한 추가 검증이 필요합니다.')
bullet('3부품 복합 결함(0.71) 개선을 위해 결함을 부품별 멀티라벨/순서형으로 푸는 접근을 검토합니다.')
bullet('NSF-MAP의 Knowledge Infusion(Process Ontology 기반 센서 범위 위반 패널티)을 향후 확장 과제로 둡니다.')
H2('6.3 역할 분담')
bullet('하현성: 멀티모달 융합 아키텍처 설계 및 코드 통합, 센서 브랜치(BiLSTM) 구현·학습, Kaggle 실행 환경 및 데이터 파이프라인 구축.')
bullet('김주헌: 이미지 브랜치(ResNet18) 구현·학습, 데이터 탐색 및 라벨 정리/클래스 병합, 평가 지표 시각화(Confusion Matrix 등).')
bullet('공동 수행: Fusion 모델 통합 학습 및 디버깅, 3모델 성능 비교 분석, 발표 자료 및 보고서 작성.')
H2('6.4 공유·소통 및 산출물 링크')
body_p('팀 작업은 Kaggle 노트북 공유와 공유 문서를 통해 진행하였습니다. 초기 실험은 하현성 계정의 Kaggle 노트북에서 시작하였으나, 무료 GPU의 주 30시간 한도 때문에 학습량이 많아지면서 김주헌 계정의 노트북으로 작업을 이전하여 최종 실험을 수행하였습니다. 따라서 작업의 흐름은 하현성 노트북(초기 버전들의 시행착오)에서 김주헌 노트북(최종 실행)으로 이어지며, 각 노트북의 버전 히스토리를 순서대로 확인하시면 전체 진행 경과를 따라가실 수 있습니다.')
bullet('최종 실행 노트북(공개): https://www.kaggle.com/code/kimjoohoen/multimodal-deep-learning-for-manufacturing-anomaly')
bullet('초기·참고 노트북(버전 히스토리): https://www.kaggle.com/code/hyunsungha/multimodal-anomaly-detection')
bullet('전처리 CSV 데이터셋(공개): https://www.kaggle.com/datasets/hyunsungha/ff-multimodal-csv')

# Word가 문서 열 때 필드(목차) 자동 갱신하도록 설정
from docx.oxml import OxmlElement
settings = doc.settings.element
uf = settings.find(qn('w:updateFields'))
if uf is None:
    uf = OxmlElement('w:updateFields'); settings.insert(0, uf)
uf.set(qn('w:val'), 'true')

doc.save(OUT)
print('저장:', OUT)
print('단락:', len(doc.paragraphs), '표:', len(doc.tables), '이미지:', len(doc.inline_shapes))
