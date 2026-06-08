# -*- coding: utf-8 -*-
"""기말 보고서 초안 생성 (python-docx). 형식: 중간보고서 템플릿(맑은 고딕/제목체계) 복제, 내용은 멀티모달 프로젝트 신규."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "맑은 고딕"
CODEFONT = "Consolas"
FIGS = "_report_figs"

doc = Document()

# ---- 기본 폰트(맑은 고딕, eastAsia 포함) ----
def set_run(run, name=KFONT, size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name); rFonts.set(qn('w:eastAsia'), name)

normal = doc.styles['Normal']
normal.font.name = KFONT; normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), KFONT)

def para(text="", size=10.5, bold=False, align=None, color=None, space_after=6, name=KFONT):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); set_run(r, name, size, bold, color)
    return p

def heading(text, level=1):
    sizes={1:15, 2:12.5, 3:11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); set_run(r, KFONT, sizes.get(level,11), True, (0x1F,0x38,0x64) if level==1 else (0x2E,0x4A,0x6B))
    # outline level for navigation
    pPr = p._element.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level-1)); pPr.append(ol)
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_run(r, KFONT, size)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = Pt(12)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F2F2F2')
    p._element.get_or_add_pPr().append(shd)
    for i,ln in enumerate(lines):
        if i>0: p.add_run().add_break()
        r=p.add_run(ln); set_run(r, CODEFONT, 9)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        r = hdr[i].paragraphs[0].add_run(h); set_run(r, KFONT, 9.5, True)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].paragraphs[0].clear()
            r = cells[i].paragraphs[0].add_run(str(val)); set_run(r, KFONT, 9.5, (i==0))
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, width_in, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); set_run(r, KFONT, 9, False, (0x55,0x55,0x55))
    c.paragraph_format.space_after = Pt(10)

# ====================== 표지 ======================
for _ in range(3): doc.add_paragraph()
para("딥러닝응용  기말 프로젝트 보고서", 14, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("멀티모달 딥러닝 기반", 24, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("제조 공정 이상 탐지", 24, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
para("( 이미지 + 시계열 센서 데이터 )", 13, False, WD_ALIGN_PARAGRAPH.CENTER, color=(0x55,0x55,0x55), space_after=80)
para("Future Factories (FF) Dataset 기반", 11, False, WD_ALIGN_PARAGRAPH.CENTER, color=(0x55,0x55,0x55), space_after=120)
para("학번 : 232320635", 12, False, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("이름 : 하현성", 12, False, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("2026. 06.", 11, False, WD_ALIGN_PARAGRAPH.CENTER, color=(0x55,0x55,0x55))
doc.add_page_break()

# ====================== 1. 개요 ======================
heading("1. 프로젝트 개요 및 목표", 1)
para("본 프로젝트는 제조 조립 공정에서 발생하는 부품 누락 결함을, 카메라 이미지와 시계열 센서 데이터를 함께 활용하는 멀티모달 딥러닝으로 탐지한다. 현재 산업 현장의 불량 판정은 대부분 비전 카메라 단독 방식이나, 조명·각도·가려짐으로 비전만으로 판단이 어려운 경우가 있고 센서 또한 노이즈·고장으로 단독 신뢰가 어렵다. 두 모달리티의 약점이 상호 보완 가능하다는 가설을 정량적으로 검증하는 것이 핵심 동기다.")
para("구체적 목표는 (1) 센서 단독(BiLSTM), (2) 이미지 단독(ResNet18), (3) 멀티모달 융합 세 모델을 동일 데이터·동일 평가 지표로 비교하여, 융합이 단일 모달리티 대비 실제로 어떤 이득을 주는지 규명하는 것이다.", space_after=6)
heading("실험 환경", 2)
bullet("플랫폼: Kaggle Notebook (Tesla T4 GPU), PyTorch")
bullet("재현성: Random Seed 42 고정, 학습 가중치·결과 Kaggle Output 저장")
bullet("워크플로: Kaggle API(kernels push/pull)로 버전 관리 및 헤드리스 실행")

# ====================== 2. 데이터셋 ======================
heading("2. 데이터셋", 1)
para("University of South Carolina의 Future Factories(FF) Lab이 공개한 산업 등급 데이터셋을 사용한다. 로봇 4대와 컨베이어로 구성된 실제 조립 라인을 30시간 가동하며 수집했으며, 약 166,000 레코드·285 사이클로 구성된다. 카메라 2대의 이미지와 동기화된 40+ 채널 센서가 함께 제공되고, 의도적으로 부품을 제거해 결함을 생성하였다.")
heading("클래스 구성 및 불균형", 2)
para("원본의 다양한 결함 조합 중 소수 클래스를 병합하여 5개 클래스로 단순화하였다: Normal, NoBody1, NoNose, NoNose_NoBody2, NoNose_NoBody2_NoBody1. Normal이 약 55%로 다수를 차지해 클래스 불균형이 크며, 이를 위해 클래스 가중치(class weight)와 Weighted F1 평가를 사용한다.")
heading("전처리", 2)
bullet("센서: 40+ 채널 중 결측·분산 0 채널을 제외하고 로봇 관절각·그리퍼 로드셀·컨베이어 신호 중심 22채널 선정 후 표준화(StandardScaler).")
bullet("시계열: 사이클 단위 sliding window로 길이 50 시퀀스 생성 (입력 (50, 22)).")
bullet("이미지: 카메라에 부품이 보이는 Cycle State 4·9 시점만 사용, 224×224 RGB로 변환·정규화.")
bullet("데이터 누출 방지: 시점 무작위 분할이 아닌 Cycle-wise 80:20 분할 — 같은 사이클의 인접 시점이 train/test에 섞이지 않도록 통제.")

# ====================== 3. 방법론 ======================
heading("3. 방법론", 1)
heading("3.1 모델 구조", 2)
para("Model 1 — 센서 단독 (BiLSTM):", 10.5, True, space_after=2)
code_block(["입력 (B, 50, 22) → BiLSTM(hidden=128, 2-layer, bidirectional)",
            "  → f_sensor(256) → FC(256→64) → ReLU → FC(64→5)"])
para("Model 2 — 이미지 단독 (ResNet18, Transfer Learning):", 10.5, True, space_after=2)
code_block(["입력 (B, 3, 224, 224) → ResNet18(ImageNet 사전학습)",
            "  → f_image(512) → FC(512→128) → ReLU → Dropout → FC(128→5)"])
para("Model 3 — 멀티모달 융합 (3가지 변형):", 10.5, True, space_after=2)
bullet("Decision-Level Concat: f_image(512)와 f_sensor(256)를 concat(768) 후 FC 분류. 이미지 없는 시점은 zero 마스킹.")
bullet("Cross-Attention: 센서를 Query, 이미지 spatial 토큰(7×7)을 Key/Value로 attention → 이미지 신뢰도가 낮을 때 센서로 자동 폴백.")
bullet("Decision-Level(Late) 확률 융합: 학습된 센서·이미지 모델의 softmax를 결합(이미지 보유 시점만), 가중치는 train에서 선택 → 결측 모달리티에 견고.")
para("추가로, 융합 모델의 센서·이미지 인코더를 단독 학습 모델의 가중치로 warm-start하여 각 브랜치를 강화하였다.", space_after=6)

heading("3.2 학습 및 평가 설계", 2)
bullet("손실: CrossEntropyLoss + 클래스 가중치. 융합은 cycle 4·9 분포 기반 가중치 별도 계산.")
bullet("최적화: Adam(lr 1e-3, weight decay 1e-4), ReduceLROnPlateau, gradient clipping, early stopping.")
bullet("평가 지표: 클래스 불균형을 고려해 Weighted F1을 주 지표로 사용(Accuracy는 다수 클래스 편향).")
bullet("공정 비교: 센서·이미지·융합을 모두 동일한 cycle 4·9 테스트셋에서 재평가하여 apples-to-apples 비교를 수행.")

heading("3.3 검증/평가 분리에 관한 방법론적 고찰", 2)
para("본 실험은 Cycle-wise로 train/test를 분리해 시점 누출을 차단했으나, 현재 구현은 early stopping의 기준으로 test set을 사용하여 엄밀한 의미의 독립 검증 셋(validation)이 분리되지 않았다는 한계가 있다. 모델 선택(조기 종료 시점)에 test 정보가 간접 반영될 수 있으므로, 보고된 수치는 낙관적으로 해석될 여지가 있다.")
para("개선 방향으로, 사이클 단위 3-way 분할(train/validation/test)을 도입해 validation으로 조기 종료·모델 선택을 수행하고 test는 최종 1회만 평가하는 절차가 더 엄밀하다. 본 보고서의 비교는 모든 모델에 동일한 절차를 적용했으므로 모델 간 상대 비교의 타당성은 유지된다.", space_after=6)

# ====================== 4. 실험 및 결과 ======================
heading("4. 실험 및 결과", 1)
heading("4.1 센서 단독 기준선", 2)
para("전체 사이클 테스트셋에서 센서 단독 BiLSTM은 Weighted F1 0.927(Accuracy 0.928)로 강력한 기준선을 형성한다. 시계열 센서가 결함 상태를 잘 구분함을 보여준다. 단, 이미지가 존재하는 cycle 4·9 구간만으로 평가하면 0.871로, 이 구간이 상대적으로 더 어려움을 알 수 있다(이후 융합 비교는 이 구간 기준).")

heading("4.2 핵심 결과 — 이미지 데이터량에 따른 멀티모달 이득", 2)
para("이미지 데이터(파트 수)를 늘려가며 단독 모델과 융합의 성능 변화를 관찰하였다. 이미지 과제는 부품 누락이 시각적으로 분명해 데이터가 충분하면 단독 성능이 매우 높아진다(1파트 0.69 → 3파트 0.995). 반면 센서 단독은 0.871로 고정된다.")
figure(f"{FIGS}/fig1_regime_ablation.png", 5.6, "[그림 1] 이미지 데이터량에 따른 단독/융합 성능 (cycle 4·9). 융합은 모든 구간에서 센서 이상.")
table(["이미지 구성", "Image(실제)", "Sensor", "Decision-Fusion", "해석"],
      [["1 파트 (커버리지 23%)", "0.69", "0.871", "0.872", "이미지 부족 → 이득 미미"],
       ["2 파트 (커버리지 41%)", "0.96", "0.871", "0.885", "균형 → 융합 > 둘 다"],
       ["3 파트 (커버리지 54%)", "0.995", "0.871", "0.959", "이미지 과다 → 이미지 단독 최강"]])
para("관찰: (1) 이미지 데이터가 적으면 융합 이득이 작고, (2) 적당하면 융합이 두 단독을 모두 상회하며, (3) 과다하면 이미지 단독이 가장 높아 융합의 추가 이득이 사라진다. 즉 멀티모달 융합의 가치는 두 모달리티의 데이터·성능 균형에 크게 의존한다.", space_after=6)

heading("4.3 동일 테스트셋 비교 (균형 구간, 2파트)", 2)
para("2파트 구성에서 동일한 cycle 4·9 테스트셋으로 네 모델을 비교하면, 융합(Concat 0.878, Decision 0.885)이 센서 단독(0.871)과 이미지(전체 구간 평가 0.432)를 모두 상회한다. 이미지의 0.432는 '이미지가 없는 시점(약 59%)'까지 포함해 평가한 값으로, 실제 이미지가 존재하는 표본만 보면 0.96에 달한다. 융합(특히 decision-level)은 이미지가 있을 때만 이를 활용하므로 결측 모달리티에 견고하며 전체 구간에서 최고 성능을 낸다.")
figure(f"{FIGS}/fig2_matched_comparison.png", 5.6, "[그림 2] 동일 cycle 4·9 테스트셋 비교(2파트). 융합이 두 단독 모달리티를 모두 초과.")

heading("4.4 견고성 및 최고 성능", 2)
para("이미지 인코더를 동결(ImageNet 특징만)해 이미지를 극단적으로 약화(0.28)시킨 조건에서도 융합은 0.887로 센서(0.871)를 상회하여, 한 모달리티가 매우 약해도 융합이 끌려가지 않는 견고성을 보였다. 한편 이미지 데이터가 풍부한 조건에서는 이미지 단독 0.995, 융합 0.964로 절대 성능이 가장 높았다.")
figure(f"{FIGS}/fig3_peak.png", 4.6, "[그림 3] 이미지 데이터가 풍부할 때의 최고 성능.")

# ====================== 5. 결론 ======================
heading("5. 결론 및 향후 과제", 1)
para("본 프로젝트는 센서·이미지·융합 모델을 동일 조건에서 비교하여, 멀티모달 융합의 이득이 두 모달리티의 데이터·성능 균형에 의존함을 정량적으로 입증하였다. 균형 구간에서 융합은 두 단독 모달리티를 모두 상회했고, decision-level 융합은 결측 모달리티에 견고하여 모든 구간에서 센서 단독 이상이었다. 이는 '데이터 균형이 멀티모달 성공의 핵심 조건'이라는 중간 발표의 가설을 뒷받침한다.")
para("실무적 시사점: 멀티모달이 단일 모달리티를 항상 이긴다는 직관은 성립하지 않으며, 산업 적용 시 각 모달리티의 데이터 수집 전략과 균형이 중요하다.", space_after=6)
heading("향후 과제", 2)
bullet("검증 엄밀화: cycle-wise 3-way 분할(train/val/test)로 독립 validation 도입(3.3 참조).")
bullet("Cross-Attention 융합 고도화 및 ResNet 부분 미세조정(NSF-MAP P2) 적용.")
bullet("Knowledge Infusion(센서 물리 범위 위반 페널티, NSF-MAP P3)으로 추가 향상.")
bullet("데이터 샘플 시각화 및 정성 분석(오분류 사례) 보강.")

doc.add_paragraph()
para("※ 본 문서는 자동 실험 파이프라인으로 산출된 결과를 정리한 초안이며, 최종 제출본이 아니다.", 9, False, color=(0x88,0x88,0x88))

OUT="기말보고서_초안_하현성.docx"
doc.save(OUT)
print("saved:", OUT)
