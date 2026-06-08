"""중간 발표 스크립트 + 모델 구조 Q&A 대응 Word 문서 (컴팩트 버전)
- Page 1: Slide 1~7 (목표·관련연구·데이터셋·아키텍처·학습전략)
- Page 2: Slide 8~12 (실행환경·결과·개선·스케줄·QA)
- Page 3+: Q&A 대응 자료
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\Hyunsung Ha\Desktop\아주대학교\4학년\딥러닝응용\기말 멀티모달 프로젝트\딥러닝응용_3조_중간발표_스크립트.docx"

# ── 색상 ──────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x16, 0xA2, 0x95)
CORAL = RGBColor(0xE8, 0x58, 0x50)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(9.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 좁은 여백
for section in doc.sections:
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_para(text, *, size=9.5, bold=False, italic=False, color=DARK, align=None,
             before=0, after=2, indent_left=None, line=1.25, font='맑은 고딕'):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if indent_left is not None:
        fmt.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return p


def add_h1(text):
    add_para(text, size=14, bold=True, color=NAVY, before=4, after=4)


def add_h2(text, color=NAVY, size=11):
    add_para(text, size=size, bold=True, color=color, before=4, after=2)


def add_h3(text):
    add_para(text, size=10, bold=True, color=TEAL, before=3, after=1)


def add_script(text):
    """슬라이드 발표 스크립트 한 단락"""
    add_para(text, size=9.5, color=DARK, indent_left=0.2, after=2, line=1.3)


def add_quote(text):
    p = add_para(text, size=9, italic=True, color=GRAY, indent_left=0.3, after=4, line=1.3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '14')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '16A295')
    pBdr.append(left)
    pPr.append(pBdr)


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_qna(question, answer):
    add_para("Q. " + question, size=9.5, bold=True, color=NAVY, after=1)
    add_quote(answer)


def add_table_with_header(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "1B2A4A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    for ri, row_data in enumerate(rows):
        tr = table.rows[ri + 1]
        is_zebra = ri % 2 == 1
        for ci, item in enumerate(row_data):
            cell = tr.cells[ci]
            if isinstance(item, tuple):
                text, bold, color = item
            else:
                text, bold, color = item, False, DARK
            if is_zebra:
                set_cell_shading(cell, "F0F4F8")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


# ═══════════════════════════════════════════════════════════════
# PAGE 1: Slide 1 ~ 7
# ═══════════════════════════════════════════════════════════════
add_h1("PART 1. 발표 스크립트 (Slide 1~7)")

add_h2("Slide 1 — 표지")
add_script("안녕하세요. 3조 김주헌, 하현성입니다. 저희가 진행 중인 멀티모달 딥러닝 기반 제조 공정 이상 탐지 프로젝트의 중간 진행 상황을 발표드리겠습니다.")

add_h2("Slide 2 — 프로젝트 동기")
add_script("프로젝트의 출발점은 저희 팀의 현장 경험입니다. 팀원 전원이 자동화 장비 업계에서 근무하면서 제조 공정에서 딥러닝이 어떻게 활용되는지 직접 접하고 있습니다.")
add_script("현재 산업 현장의 불량 판정은 대부분 비전 카메라 단독 방식입니다. 그러나 실제로는 조명, 각도, 가려짐 때문에 비전만으로 판단이 어려운 경우가 자주 발생합니다. 반대로 센서 데이터도 노이즈나 고장 때문에 단독으로는 신뢰하기 어려운 경우가 존재합니다.")
add_script("두 모달리티의 약점이 서로 보완 가능하다는 점에 주목해서, 이미지와 센서를 함께 활용하는 멀티모달 접근이 실제로 어떤 성능 차이를 만드는지 정량적으로 확인하는 것이 본 프로젝트의 핵심 동기입니다.")

add_h2("Slide 3 — 프로젝트 목표")
add_script("구체적인 목표는 세 가지 모델을 동일한 데이터와 평가 지표로 비교하는 것입니다.")
add_script("첫 번째는 BiLSTM(바이 엘에스티엠) 기반 센서 단독 모델로, 22채널 시계열 센서를 입력으로 사용하며 전체 사이클에 적용 가능합니다.")
add_script("두 번째는 ResNet18(레즈넷 에이틴) 기반 이미지 단독 모델로, Transfer Learning을 적용하며 카메라에 부품이 보이는 Cycle 4·9 구간만 학습합니다.")
add_script("세 번째는 앞의 두 브랜치를 Decision-Level Fusion으로 결합한 멀티모달 모델입니다.")
add_script("단일 모달리티 대비 융합 모델이 어떤 보완 효과를 보이는지가 가장 중요한 관찰 포인트입니다.")

add_h2("Slide 4 — 관련 연구 (NSF-MAP)")
add_script("주요 참고 논문은 IJCAI 2025의 NSF-MAP입니다. 동일한 Future Factories 데이터셋을 사용했고, 단계적 발전 방식이 핵심입니다.")
add_script("P1 기본 Decision-Level Fusion이 72%, P2에서 Transfer Learning을 더해 88%, P3에서 Process Ontology 기반 Knowledge Infusion으로 93%를 달성합니다.")
add_script("본 프로젝트는 강의에서 학습한 LSTM·ResNet으로 모델을 재구성하고, 모달리티별 효과 비교 자체에 집중합니다. Knowledge Infusion은 향후 확장 과제로 둡니다.")

add_h2("Slide 5 — 데이터셋")
add_script("데이터셋은 사우스캐롤라이나 대학교에서 공개한 Future Factories Lab Dataset입니다. 로봇 4대와 컨베이어로 구성된 실제 조립 라인을 30시간 가동하면서 수집한 데이터로, 16만 개 레코드와 285 사이클이 있습니다.")
add_script("센서는 40채널 이상이고, 카메라 2대가 동기화 촬영합니다.")
add_script("라벨 분포를 보시면 Normal이 약 55%로 다수를 차지하고, 소수 클래스가 여러 개 있어서 클래스 불균형이 큰 문제입니다. 저희는 소수 클래스 일부를 병합해 5클래스로 단순화하고, Cycle-wise 분할로 데이터 누출을 방지했습니다.")

add_h2("Slide 6 — 전체 아키텍처")
add_script("전체 아키텍처를 흐름도로 정리하면 이렇습니다.")
add_script("센서 데이터는 22채널 50 스텝 시퀀스로 만들어 BiLSTM에 입력하고, 양방향 hidden state를 합쳐 256차원의 센서 특징 벡터를 얻습니다.")
add_script("이미지 데이터는 224×224 RGB 이미지로 변환해서 ImageNet 사전학습 ResNet18에 통과시키고 512차원 특징 벡터를 추출합니다. 이때 ResNet18은 사전학습 가중치를 그대로 사용하기 위해 backbone을 동결했습니다.")
add_script("마지막으로 두 특징을 Concatenate해서 768차원으로 만든 뒤 FC와 Dropout 레이어를 거쳐 5클래스 분류를 수행합니다. 이미지가 없는 시점에는 센서 정보만 활용하도록 마스킹 처리도 적용했습니다.")

add_h2("Slide 7 — 모델별 학습 전략")
add_script("각 모델의 학습 전략은 슬라이드의 카드에 정리했습니다. 세 모델 모두 동일한 5클래스, Cycle-wise 80대 20 분할, 동일한 평가 지표를 사용해서 오로지 모달리티 차이만 비교할 수 있도록 통제했습니다.")
add_script("이미지와 Fusion 모델은 데이터가 Cycle 4와 9 구간으로 한정되어 더 작은 학습셋을 가진다는 점이 한계입니다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PAGE 2: Slide 8 ~ 12
# ═══════════════════════════════════════════════════════════════
add_h1("PART 1. 발표 스크립트 (Slide 8~12)")

add_h2("Slide 8 — 실행 환경 및 데이터 제약")
add_script("실행 환경에서 가장 큰 제약은 데이터 용량이었습니다. 이미지 원본이 6 파트 580GB로 로컬 PC 다운로드가 사실상 불가능했습니다.")
add_script("또한 3개 모델을 반복적으로 학습·실험해야 하는 상황에서 로컬 자원을 장시간 점유하는 부담이 컸기 때문에, 빠른 프로토타이핑이 필요했습니다.")
add_script("이 문제를 Kaggle Notebook으로 해결했습니다. Kaggle 서버에 마운트된 원본 데이터를 다운로드 없이 직접 읽고, 무료로 제공되는 Tesla T4 GPU로 실험 사이클을 단축했습니다. 검증된 모델은 추후 로컬에서도 재현 가능합니다.")
add_script("다만 이미지는 6 파트 중 1 파트만 사용해서 약 9,500장이 매칭됐는데, 이는 NSF-MAP에서 사용한 약 1.5만 장과 비교해도 학술적으로 유의미한 규모입니다.")
add_script("추가로 Kaggle 세션이 idle 시 자동 종료되는 점을 고려해서 Random Seed를 42로 고정하고, 학습된 가중치와 결과를 Kaggle Output에 별도 저장해 재현성을 확보했습니다.")

add_h2("Slide 9 — 최종 실험 결과")
add_script("3개 모델 학습이 모두 완료되었습니다. 결과를 보시면 센서 단독 BiLSTM(바이 엘에스티엠)이 Weighted F1 0.927로 가장 높은 성능을 기록했습니다.")
add_script("이미지 단독 ResNet18(레즈넷 에이틴)은 F1 0.611로 가장 낮고, Fusion은 F1 0.882로 중간에 위치합니다. 직관적으로는 Fusion이 가장 좋아야 하는데, 결과가 반대로 나온 이유가 있습니다.")
add_script("이미지 데이터는 Cycle 4와 9 구간, 파트 1만 사용해서 약 9,500장에 불과합니다. 반면 센서 데이터는 전체 사이클 12만여 시퀀스로 훨씬 풍부합니다. 이 불균형 때문에 Fusion에서 이미지 브랜치가 정보를 더하기보다 노이즈로 작용한 것으로 분석됩니다.")
add_script("NSF-MAP 논문도 P1 기본 Fusion이 72%로 가장 낮았다가 Transfer Learning과 Knowledge Infusion을 추가하면서 93%까지 올라간 것과 같은 맥락입니다. 즉, 이번 결과 자체가 멀티모달 융합에서 데이터 균형이 얼마나 중요한지를 잘 보여주는 사례입니다.")

add_h2("Slide 10 — 한계와 향후 개선 방향")
add_script("한계의 핵심은 데이터 불균형입니다. 이미지 9,500장 대비 센서 12만 시퀀스로 Fusion에서 이미지 브랜치가 노이즈로 작용했습니다.")
add_script("개선은 두 축입니다. 데이터 측면에서는 이미지 파트를 1개에서 2~6개까지 확장하고, MixUp(믹스업)·CutMix(컷믹스) 등 강한 Augmentation, 그리고 Fusion 학습 시 두 모달리티 샘플링 비율을 맞추는 전략을 적용할 계획입니다.")
add_script("모델 측면에서 가장 우선순위가 높은 건 단순 Concat을 Cross-Attention으로 바꾸는 것입니다. 이미지 신뢰도가 낮을 때 센서에 자동으로 가중치가 실립니다. 추가로 ResNet18 일부 unfreeze는 NSF-MAP P2가 72%를 88%로 끌어올린 방식이고, Knowledge Infusion은 P3가 93% 달성한 핵심 기법입니다.")
add_script("이 개선들을 적용하면 Fusion이 센서 단독을 상회할 것으로 기대합니다.")

add_h2("Slide 11 — 스케줄 및 역할 분담")
add_script("스케줄은 1-3주차 데이터 확보와 모델 구현, 4-5주차에 3개 모델 학습과 비교 분석까지 완료했습니다. 남은 6주차는 추가 개선 실험과 최종 보고서 작성입니다.")
add_script("역할은 하현성이 멀티모달 융합 아키텍처와 센서 브랜치, Kaggle 환경 구축을, 김주헌이 이미지 브랜치와 데이터 탐색 및 라벨 정리, 평가 시각화를 담당했습니다. Fusion 통합과 비교 분석, 보고서 작성은 공동 수행입니다.")

add_h2("Slide 12 — Q&A")
add_script("이상입니다. 질문 받겠습니다. 감사합니다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PAGE 3+: Q&A 대응 자료
# ═══════════════════════════════════════════════════════════════
add_h1("PART 2. Q&A 대응 자료")

# ─── 1. 분류 클래스 ────────────────────────────────────
add_h2("1. 분류 클래스 구성")
add_para("원본 7가지 결함 조합 중 'Body 계열만 누락'인 3개를 NoBody1로 통합 → 최종 5클래스.",
         italic=True, color=GRAY, after=4)
add_table_with_header(
    ["원본 라벨 (key)", "최종 클래스 (value)"],
    [
        [("Normal", False, DARK), ("Normal", True, TEAL)],
        [("NoBody1", False, DARK), ("NoBody1", True, TEAL)],
        [("NoBody2", False, DARK), ("NoBody1 ← 병합", True, CORAL)],
        [("NoBody2,NoBody1", False, DARK), ("NoBody1 ← 병합", True, CORAL)],
        [("NoNose", False, DARK), ("NoNose", True, TEAL)],
        [("NoNose,NoBody2", False, DARK), ("NoNose_NoBody2", True, TEAL)],
        [("NoNose,NoBody2,NoBody1", False, DARK), ("NoNose_NoBody2_NoBody1", True, TEAL)],
    ],
    col_widths_cm=[8.5, 8.5]
)
add_table_with_header(
    ["최종 클래스", "의미", "개수", "비율"],
    [
        [("Normal", True, DARK), "정상", ("90,775", False, DARK), ("54.7%", False, DARK)],
        [("NoBody1 (병합)", True, DARK), "Body 계열 결함", ("4,016", False, DARK), ("2.4%", False, DARK)],
        [("NoNose", True, DARK), "Nose 누락", ("19,307", False, DARK), ("11.7%", False, DARK)],
        [("NoNose_NoBody2", True, DARK), "Nose+Body2 누락", ("25,206", False, DARK), ("15.2%", False, DARK)],
        [("NoNose_NoBody2_NoBody1", True, DARK), "3개 모두 누락", ("26,628", False, DARK), ("16.0%", False, DARK)],
    ],
    col_widths_cm=[6.0, 4.5, 3.0, 3.5]
)

add_divider()

# ─── 2. Model 1 ─────────────────────────────────────────
add_h2("2. Model 1 — BiLSTM (센서 단독)")
add_h3("아키텍처")
add_code("입력 (batch, 50, 22) → BiLSTM (hidden=128, 2-layer, bidirectional)")
add_code("  → f_sensor (256) → FC(256→128) → Dropout(0.3) → FC(128→5)")

add_h3("핵심 하이퍼파라미터")
add_table_with_header(
    ["항목", "값", "근거"],
    [
        ["seq_len", ("50 step", True, DARK), "~10Hz 샘플링 기준 5초 윈도우"],
        ["hidden_dim", ("128", True, DARK), "NSF-MAP 및 강의 노트 표준 설정"],
        ["num_layers", ("2", True, DARK), "1층 부족, 3층 gradient 불안정"],
        ["dropout", ("0.2 + 0.3", True, DARK), "LSTM 내부 + FC 직전, 과적합 방지"],
        ["학습 샘플", ("121,550", True, DARK), "Sliding Window 적용 후"],
    ],
    col_widths_cm=[3.0, 2.5, 11.0]
)

add_h3("예상 질문")
add_qna("왜 단방향 LSTM 아닌 BiLSTM?",
        "제조 공정 사이클은 사후 분석이 가능한 task라 미래 정보 활용 가능합니다. 양방향 구조가 anomaly의 전조와 잔여 패턴을 모두 포착할 수 있고, 강의 10-2 p13에 양방향 모델이 명시되어 있어 강의 범위 내 선택입니다.")
add_qna("왜 Transformer 아닌 LSTM?",
        "데이터가 12만 시퀀스로 Transformer엔 비효율적이고, 시퀀스 길이 50으로 짧아 LSTM이 충분합니다. 강의 범위 내 모델 사용 원칙도 있습니다.")
add_qna("22채널 어떻게 선정?",
        "원본 40+ 채널 중 결측 많거나 분산 0인 채널 제외하고, 핵심 로봇 관절각도·그리퍼 로드셀·컨베이어 신호 중심으로 22개를 선정했습니다.")

add_divider()

# ─── 3. Model 2 ─────────────────────────────────────────
add_h2("3. Model 2 — ResNet18 (이미지 단독)")
add_h3("아키텍처")
add_code("입력 (batch, 3, 224, 224) → ResNet18 (ImageNet pretrained, frozen)")
add_code("  → f_image (512) → FC(512→128) → Dropout(0.5) → FC(128→5)")

add_h3("핵심 하이퍼파라미터")
add_table_with_header(
    ["항목", "값", "근거"],
    [
        ["backbone", ("ResNet18", True, DARK), "강의 7-2 p16, 9,500장에 적합"],
        ["Pretrained", ("ImageNet", True, DARK), "Transfer Learning, 일반 시각 특징 활용"],
        ["Freeze", ("True", True, DARK), "9,500장으로 fine-tune 시 과적합 위험"],
        ["Augmentation", ("Flip/Rot/Color", True, DARK), "HFlip + Rotation(10°) + ColorJitter(0.2)"],
        ["학습 샘플", ("9,542장", True, DARK), "Cycle 4·9 한정, Part 1만"],
    ],
    col_widths_cm=[3.0, 2.5, 11.0]
)

add_h3("예상 질문")
add_qna("왜 ResNet50/EfficientNet 아닌 ResNet18?",
        "9,500장은 deeper model엔 부족하여 작은 모델이 유리합니다. 강의 7-2에서 ResNet을 다뤘기에 강의 범위 내 선택이며, NSF-MAP의 EfficientNet-B0를 단순화한 것입니다.")
add_qna("Backbone 동결 이유?",
        "ImageNet의 일반 시각 특징이 산업 이미지에도 유효하고, 9,500장으로 전체 fine-tune 시 과적합 위험이 큽니다. 향후 개선에 일부 unfreeze 포함되어 있습니다.")
add_qna("왜 Cycle 4·9만 사용?",
        "그 외 사이클은 카메라에 부품이 보이지 않습니다(컨베이어 이동 중). 결함 판별에 의미 있는 시각 정보가 Cycle 4·9에만 존재합니다.")
add_qna("F1 0.611이 낮은 이유?",
        "9,500장에 5클래스 불균형이 가장 큰 원인입니다. Precision 0.53이 Recall 0.73보다 낮은데, 다수 클래스 편향 학습 신호입니다. 클래스 가중치 재계산으로 일부 개선했지만 절대 데이터 양 부족이 한계입니다.")

add_divider()

# ─── 4. Model 3 ─────────────────────────────────────────
add_h2("4. Model 3 — Decision-Level Fusion (멀티모달)")
add_h3("아키텍처")
add_code("센서 → BiLSTM → f_sensor(256)  ─┐")
add_code("                                  ├─ Concat → 768 → FC(768→256) → Dropout → FC(256→5)")
add_code("이미지 → ResNet18(frozen) → f_image(512) ─┘")

add_h3("핵심 설계")
add_bullet("Decision-Level Fusion (Late Fusion): 각 모달리티 특징 추출 후 concat")
add_bullet("마스킹: 이미지 없는 시점은 zero 벡터 → 사실상 센서 정보만으로 분류")
add_bullet("학습: BiLSTM learnable, ResNet18 frozen (계산 효율)")
add_bullet("Loss: CrossEntropyLoss + 이미지 분포 기반 Class Weight")

add_h3("예상 질문")
add_qna("Early/Feature-level Fusion은 안 했나?",
        "Early Fusion은 raw 데이터 차원이 너무 달라(시퀀스 vs 이미지) 결합이 어렵습니다. Decision-Level은 NSF-MAP P1 baseline과 동일해 직접 비교가 용이합니다.")
add_qna("Concat 외에 다른 융합 방식은?",
        "본 발표는 baseline 비교가 목적입니다. 슬라이드 10 개선안에 Cross-Attention을 명시했습니다. 적용 시 이미지 신뢰도가 낮으면 자동으로 센서에 가중치가 부여됩니다.")
add_qna("Fusion이 센서 단독보다 낮은 이유?",
        "이미지 9,500장 vs 센서 12만 시퀀스의 데이터 불균형이 원인입니다. ResNet18 출력이 불안정하여 768차원 concat에서 이미지 정보가 노이즈로 작용했습니다. NSF-MAP P1이 72%로 가장 낮았던 것과 동일한 현상입니다.")
add_qna("멀티모달 시도가 실패한 것 아닌가?",
        "아닙니다. 데이터 균형이 멀티모달 성공의 핵심임을 입증한 baseline 결과입니다. 개선 방향 6가지는 모두 선행 연구로 검증된 방법이며, 적용 시 Fusion이 센서 단독을 상회할 가능성이 충분합니다.")

add_divider()

# ─── 5. 공통 평가 조건 ───────────────────────────────────
add_h2("5. 공통 평가 조건")
add_table_with_header(
    ["항목", "값 / 설명"],
    [
        [("Train/Test", True, DARK), "Cycle-wise 80:20 (228 학습 / 57 평가, 데이터 누출 방지)"],
        [("평가 지표", True, DARK), "Accuracy, Weighted F1 (메인), Precision, Recall"],
        [("클래스", True, DARK), "5개 (Normal + 4 결함, 소수 클래스 일부 병합)"],
        [("Random Seed", True, DARK), "42 (numpy, torch, random)"],
        [("학습 환경", True, DARK), "Kaggle Notebook + Tesla T4 + PyTorch"],
    ],
    col_widths_cm=[3.5, 13.0]
)
add_qna("왜 Weighted F1이 메인 지표?",
        "Accuracy는 다수 클래스(Normal 55%) 편향으로 불균형 데이터에 부적절. Weighted F1은 클래스 빈도 가중 평균이라 적합합니다.")
add_qna("Cycle-wise 분할이 왜 중요?",
        "시점별 무작위 분할 시 같은 사이클의 인접 시점이 Train/Test 양쪽에 들어가 정보 누출이 발생합니다. Cycle 단위 분할이 실전 일반화 성능을 정확히 측정합니다.")

add_divider()

# ─── 6. 최종 결과 ─────────────────────────────────────────
add_h2("6. 최종 결과 & NSF-MAP 비교")
add_table_with_header(
    ["모델", "Accuracy", "Precision", "Recall", "F1"],
    [
        [("M1 · Sensor (BiLSTM)", True, TEAL), ("0.9276", True, DARK), "0.9276", "0.9276", ("0.9267", True, TEAL)],
        [("M2 · Image (ResNet18)", True, CORAL), "0.7263", "0.5276", "0.7263", ("0.6112", True, CORAL)],
        [("M3 · Fusion", True, DARK), "0.8941", "0.8891", "0.8941", ("0.8824", True, DARK)],
    ],
    col_widths_cm=[5.5, 2.8, 2.8, 2.8, 2.6]
)
add_qna("NSF-MAP P1 Fusion이 72%, 우리는 0.88 — 우리가 더 잘한 것?",
        "단순 1:1 비교는 어렵습니다. (1) 평가 데이터 구성이 다릅니다 — 저희는 Cycle 4·9 시점만, NSF-MAP은 전체 사이클. (2) 저희 센서 브랜치가 F1 0.927로 매우 강해 Fusion에서도 그 정보가 그대로 활용됩니다. (3) NSF-MAP P1은 baseline 성격이라 ImageNet 사전학습 같은 강화 기법을 제외, 저희는 P1부터 적용. '다른 조건에서 측정한 다른 baseline'으로 보는 것이 정확합니다.")

add_divider()

# ─── 7. 개선 방향 ────────────────────────────────────────
add_h2("7. 향후 개선 방향")
add_para("데이터 측면", bold=True, size=10, color=NAVY, before=2, after=1)
add_bullet("① 이미지 파트 확장 (1 → 2~6): 9,542장 → 약 53,000장 (~5.5배)")
add_bullet("② MixUp/CutMix Augmentation: 가림 학습 + 가상 샘플 생성")
add_bullet("③ 모달리티 비율 맞춤 샘플링: 센서 다운샘플 또는 이미지 업샘플")
add_para("모델 측면", bold=True, size=10, color=CORAL, before=2, after=1)
add_bullet("Ⓐ Cross-Attention Fusion (최우선): Q=f_sensor, K=V=f_image, 이미지 노이즈 시 자동 센서 우회")
add_bullet("Ⓑ ResNet18 일부 unfreeze (layer4+FC): NSF-MAP P2 전략 (72→88%, +16%p)")
add_bullet("Ⓒ Knowledge Infusion: 센서 범위 위반 시 손실 함수 페널티 (NSF-MAP P3, 88→93%)")

add_divider()

# ─── 8. 기타 질문 ───────────────────────────────
add_h2("8. 기타 자주 묻는 질문")
add_qna("개선 우선순위는?",
        "Cross-Attention(Ⓐ) > 데이터 확장(①) > Transfer Learning 강화(Ⓑ) 순. Cross-Attention은 코드 수정만으로 즉시 검증 가능합니다.")
add_qna("정말 Fusion이 센서 단독을 넘을 수 있을까?",
        "NSF-MAP이 동일 데이터셋에서 72→88→93%까지 끌어올린 선례가 있어 가능성은 충분합니다. 저희는 P1부터 강화된 baseline이라 향상폭은 작아도 '단독 모달리티 초월'이 현실적 목표입니다.")
add_qna("이 프로젝트의 학술적 의의는?",
        "멀티모달 융합이 단일 모달리티를 항상 이긴다는 막연한 직관에 반해, 데이터 균형이 핵심 조건임을 정량적으로 입증한 사례입니다. 실제 산업 현장에 적용 시 데이터 수집 전략의 중요성을 시사합니다.")
add_qna("강의에서 배운 내용을 어디에 적용?",
        "강의 10-2의 RNN/LSTM/GRU와 양방향 모델 → BiLSTM 적용. 강의 7-2의 ResNet Residual Connection과 ImageNet 사전학습 → ResNet18 적용. 강의 9-1의 멀티모달 프로젝트 안내 CNN+RNN 융합 → Fusion 모델에 적용.")

doc.save(OUT)
print(f"저장 완료: {OUT}")
